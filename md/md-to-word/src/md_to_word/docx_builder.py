from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .markdown_parser import Node
from .image_handler import ImageHandler
from .mermaid_renderer import MermaidRenderer


def add_hyperlink(paragraph, text: str, url: str):
    """Add a clickable hyperlink to a python-docx paragraph."""
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)
    return hyperlink


class DocxBuilder:
    def __init__(
        self,
        image_handler: ImageHandler,
        mermaid_renderer: MermaidRenderer,
        max_image_width: float = 6.2,
    ):
        self.image_handler = image_handler
        self.mermaid_renderer = mermaid_renderer
        self.max_image_width = max_image_width
        self.document = Document()
        self._configure_styles()

    def _configure_styles(self):
        styles = self.document.styles

        normal = styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.5)

        for name, size in [
            ("Title", 24),
            ("Heading 1", 18),
            ("Heading 2", 15),
            ("Heading 3", 13),
            ("Heading 4", 11),
        ]:
            style = styles[name]
            style.font.name = "Aptos Display"
            style.font.size = Pt(size)

        if "Code Block" not in [s.name for s in styles]:
            code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = "Courier New"
            code_style.font.size = Pt(8.5)

    def build(self, nodes: list[Node], output_path: Path):
        for node in nodes:
            self._add_node(node)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(output_path)

    def _add_node(self, node: Node):
        if node.type == "heading":
            level = min(max(node.level, 1), 9)
            p = self.document.add_heading(node.content, level=level)
            return

        if node.type == "paragraph":
            p = self.document.add_paragraph()
            self._add_inline(p, node.children)
            return

        if node.type == "code":
            p = self.document.add_paragraph(style="Code Block")
            if node.attrs.get("language"):
                run = p.add_run(f"[{node.attrs['language']}]\n")
                run.bold = True
            p.add_run(node.content.rstrip())
            return

        if node.type == "mermaid":
            self._add_mermaid(node.content)
            return

        if node.type == "image":
            self._add_image(node.attrs.get("src", ""), node.attrs.get("alt", ""))
            return

        if node.type == "list":
            self._add_list(node)
            return

        if node.type == "blockquote":
            p = self.document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            run = p.add_run("❝ ")
            run.bold = True
            for child in node.children:
                if child.type == "paragraph":
                    self._add_inline(p, child.children)
            return

        if node.type == "table":
            self._add_table(node)
            return

        if node.type == "hr":
            p = self.document.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            pbdr.append(bottom)
            pPr.append(pbdr)

    def _add_inline(self, paragraph, nodes: list[Node]):
        bold = False
        italic = False
        strike = False
        link_url = None
        link_buffer = []

        for node in nodes:
            if node.type == "strong_open":
                bold = True
            elif node.type == "strong_close":
                bold = False
            elif node.type == "em_open":
                italic = True
            elif node.type == "em_close":
                italic = False
            elif node.type == "s_open":
                strike = True
            elif node.type == "s_close":
                strike = False
            elif node.type == "link_open":
                link_url = node.attrs.get("href", "")
                link_buffer = []
            elif node.type == "link_close":
                if link_url:
                    add_hyperlink(paragraph, "".join(link_buffer), link_url)
                link_url = None
                link_buffer = []
            elif node.type == "break":
                if link_url:
                    link_buffer.append("\n")
                else:
                    paragraph.add_run().add_break()
            elif node.type == "image":
                self._add_inline_image(paragraph, node.attrs.get("src", ""), node.attrs.get("alt", ""))
            else:
                text = node.content
                if link_url is not None:
                    link_buffer.append(text)
                else:
                    run = paragraph.add_run(text)
                    run.bold = bold
                    run.italic = italic
                    run.font.strike = strike
                    if node.type == "code_inline":
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)

        if link_url and link_buffer:
            add_hyperlink(paragraph, "".join(link_buffer), link_url)

    def _add_inline_image(self, paragraph, src: str, alt: str):
        path = self.image_handler.resolve(src)
        if path:
            run = paragraph.add_run()
            run.add_picture(str(path), width=Inches(min(self.max_image_width, 3.0)))
        elif alt:
            paragraph.add_run(f"[Image: {alt}]")

    def _add_image(self, src: str, alt: str):
        path = self.image_handler.resolve(src)
        if not path:
            self.document.add_paragraph(f"[Image not found: {src}]")
            return

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(path), width=Inches(self.max_image_width))
        except Exception:
            p.add_run(f"[Unable to embed image: {src}]")
        if alt:
            caption = self.document.add_paragraph(alt)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].italic = True

    def _add_mermaid(self, source: str):
        name = f"mermaid_{len(self.document.paragraphs) + 1}"
        try:
            image = self.mermaid_renderer.render(source, name)
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(image), width=Inches(self.max_image_width))
        except RuntimeError as exc:
            p = self.document.add_paragraph()
            r = p.add_run(f"[Mermaid diagram could not be rendered: {exc}]")
            r.italic = True

    def _add_list(self, node: Node, level: int = 0):
        ordered = node.attrs.get("ordered") == "true"

        for item in node.children:
            for child in item.children:
                if child.type == "paragraph":
                    style = "List Number" if ordered else "List Bullet"
                    p = self.document.add_paragraph(style=style)
                    self._add_inline(p, child.children)
                elif child.type == "list":
                    self._add_list(child, level + 1)

    def _add_table(self, node: Node):
        rows = node.children
        if not rows:
            return

        cols = max(len(row.children) for row in rows)
        table = self.document.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"

        for r_idx, row in enumerate(rows):
            for c_idx, cell_node in enumerate(row.children):
                if c_idx >= cols:
                    continue
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                self._add_inline(p, cell_node.children)

        # Make first row bold as the likely Markdown header.
        for run in table.rows[0].cells:
            for paragraph in run.paragraphs:
                for r in paragraph.runs:
                    r.bold = True
